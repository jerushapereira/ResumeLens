from pathlib import Path

from docx import Document
from pypdf import PdfReader


def is_allowed_file(filename, allowed_extensions):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed_extensions


def extract_resume_text(file_path):
    path = Path(file_path)
    extension = path.suffix.lower()

    if extension == ".pdf":
        text = _extract_pdf(path)
    elif extension == ".docx":
        text = _extract_docx(path)
    else:
        raise ValueError("Unsupported file type.")

    cleaned = " ".join(text.replace("\x00", " ").split())
    if not cleaned:
        raise ValueError("No readable text was found in the uploaded file.")
    return cleaned


def _extract_pdf(path):
    try:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except Exception as exc:
        raise ValueError("Unable to read this PDF. Try exporting it again as a text-based PDF.") from exc


def _extract_docx(path):
    try:
        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]

        # Resumes often keep details inside tables, so include table cells too.
        table_cells = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_cells.append(cell.text)

        return "\n".join(paragraphs + table_cells)
    except Exception as exc:
        raise ValueError("Unable to read this DOCX file. Please upload a valid Word document.") from exc
