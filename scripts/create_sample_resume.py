from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT = BASE_DIR / "sample_resumes" / "sample_resume.docx"


def add_heading(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(17, 34, 80)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def main():
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Priya Sharma")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(17, 34, 80)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("B.Tech Computer Science Student | priya.sharma@example.com | linkedin.com/in/priyasharma | github.com/priyasharma")

    add_heading(document, "Summary")
    document.add_paragraph(
        "Motivated computer science student interested in full stack web development and AI-enabled products. "
        "Comfortable building small applications with Python, Flask, JavaScript, HTML, CSS, SQL, and Git."
    )

    add_heading(document, "Education")
    document.add_paragraph("B.Tech in Computer Science, City Engineering College | Expected Graduation: 2027 | CGPA: 8.4/10")

    add_heading(document, "Skills")
    document.add_paragraph("Python, JavaScript, HTML, CSS, Flask, SQL, Git, GitHub, REST APIs, Bootstrap, Pandas, Communication")

    add_heading(document, "Projects")
    document.add_paragraph("AI Resume Analyzer")
    add_bullet(document, "Built a Flask web app that uploads resumes, extracts text from documents, and generates ATS feedback using an AI API.")
    add_bullet(document, "Added Chart.js dashboard cards, dark mode, and downloadable reports for a complete portfolio demo.")
    document.add_paragraph("Internship Tracker")
    add_bullet(document, "Created a responsive web app to track internship applications by status, deadline, and company.")
    add_bullet(document, "Used browser local storage and filtering to make the app easy to use without a backend.")

    add_heading(document, "Experience")
    document.add_paragraph("Web Development Intern, BrightPath Labs | June 2025 - August 2025")
    add_bullet(document, "Developed reusable UI components and fixed layout bugs across student-facing pages.")
    add_bullet(document, "Improved page load time by 18% by compressing assets and removing unused CSS.")

    add_heading(document, "Certifications")
    add_bullet(document, "Python for Everybody - Coursera")
    add_bullet(document, "Responsive Web Design - freeCodeCamp")

    OUTPUT.parent.mkdir(exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
